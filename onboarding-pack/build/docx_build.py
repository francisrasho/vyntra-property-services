#!/usr/bin/env python3
"""Build the editable Word version of the Vyntra Subcontractor Onboarding Pack."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0F, 0x17, 0x2A)
GOLD = RGBColor(0xB8, 0x93, 0x2B)
GOLD_BR = RGBColor(0xD4, 0xAF, 0x37)
BODY = RGBColor(0x33, 0x41, 0x55)
MUTED = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
# Base style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = BODY

sec = doc.sections[0]
sec.page_height = Mm(297); sec.page_width = Mm(210)
sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(1.8)
sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
sec.header_distance = Cm(1.1)
# Running header on pages 2+ (cover/first page excluded)
sec.different_first_page_header_footer = True
try:
    hp = sec.header.paragraphs[0]
    hp.add_run().add_picture('src/assets/logo-light.png', height=Cm(0.6))
    tr = hp.add_run('   Subcontractor Onboarding Pack')
    tr.font.size = Pt(8); tr.font.color.rgb = MUTED; tr.font.name = 'Calibri'
except Exception as e:
    print('header image skipped:', e)


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa'); m.append(e)
    tcPr.append(m)


def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'), 'none'); borders.append(e)
    tblPr.append(borders)


def line_borders(table, color='E2E8F0'):
    tbl = table._tbl
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4'); e.set(qn('w:color'), color)
        borders.append(e)
    tbl.tblPr.append(borders)


def run(p, text, size=10.5, bold=False, color=BODY, italic=False, font='Calibri'):
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color; r.font.name = font
    return r


def kicker(text):
    p = doc.add_paragraph(); p.space_after = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run(p, text.upper(), size=8.5, bold=True, color=GOLD)
    pf = p.paragraph_format; pf.space_before = Pt(4); pf.space_after = Pt(1)
    return p


def h1(text):
    p = doc.add_paragraph()
    run(p, text, size=22, bold=True, color=NAVY, font='Calibri Light')
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.space_before = Pt(0)
    # gold rule
    rp = doc.add_paragraph(); rr = run(rp, '_____', size=10, bold=True, color=GOLD_BR)
    rp.paragraph_format.space_after = Pt(8); rp.paragraph_format.space_before = Pt(0)
    return p


def h2(text):
    p = doc.add_paragraph()
    run(p, text, size=13, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    return p


def para(text, size=10.5, color=BODY):
    p = doc.add_paragraph(); run(p, text, size=size, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p


def checklist(items, mark='✔'):
    for it in items:
        p = doc.add_paragraph()
        run(p, mark + '  ', size=10.5, bold=True, color=GOLD)
        run(p, it, size=10)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.4)


def callout(title, text, fill='FBF7E8'):
    dark = fill.upper() in ('0F172A', '1E293B')
    border = '334155' if dark else ('FECDCA' if fill.upper() == 'FEF3F2' else ('ABEFC6' if fill.upper() == 'ECFDF3' else 'EAD9A6'))
    title_color = GOLD_BR if dark else NAVY
    text_color = RGBColor(0xCB, 0xD5, 0xE1) if dark else BODY
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    line_borders(t, border)
    c = t.cell(0, 0); shade(c, fill); set_cell_margins(c, 120, 120, 160, 160)
    c.paragraphs[0].text = ''
    if title:
        p = c.paragraphs[0]; run(p, title, size=10.5, bold=True, color=title_color)
        p.paragraph_format.space_after = Pt(2)
        p2 = c.add_paragraph(); run(p2, text, size=9.5, color=text_color)
    else:
        p = c.paragraphs[0]; run(p, text, size=9.5, color=text_color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def kv_table(rows, head=('Item', 'Detail'), widths=(0.38, 0.62)):
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    line_borders(t)
    hdr = t.rows[0].cells
    for i, htxt in enumerate(head):
        shade(hdr[i], '0F172A'); set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]; run(p, htxt, size=9, bold=True, color=WHITE)
    for ri, (a, b) in enumerate(rows):
        cells = t.add_row().cells
        if ri % 2 == 1:
            shade(cells[0], 'F8FAFC'); shade(cells[1], 'F8FAFC')
        for ci, txt in enumerate((a, b)):
            set_cell_margins(cells[ci])
            p = cells[ci].paragraphs[0]
            run(p, txt, size=9.5, bold=(ci == 0), color=(NAVY if ci == 0 else BODY))
    # column widths
    for row in t.rows:
        row.cells[0].width = Cm(6.4); row.cells[1].width = Cm(10.6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def page_break():
    doc.add_page_break()


# ============ COVER ============
for _ in range(2):
    doc.add_paragraph()
band = doc.add_table(rows=1, cols=1); band.alignment = WD_TABLE_ALIGNMENT.CENTER
no_borders(band)
c = band.cell(0, 0); shade(c, '0F172A'); set_cell_margins(c, 500, 500, 360, 360)
p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.add_run().add_picture('src/assets/logo-on-navy.png', width=Cm(9.0))
p.paragraph_format.space_after = Pt(14)
p2 = c.add_paragraph(); run(p2, 'Subcontractor', size=34, bold=True, color=WHITE, font='Calibri Light')
p2.paragraph_format.space_after = Pt(0)
p3 = c.add_paragraph(); run(p3, 'Onboarding Pack', size=34, bold=True, color=GOLD_BR, font='Calibri Light')
p3.paragraph_format.space_after = Pt(10)
p4 = c.add_paragraph()
run(p4, 'Everything you need to deliver premium property maintenance & cleaning with Vyntra — our standards, your payments, and how we work together.', size=11, color=RGBColor(0x94, 0xA3, 0xB8))
doc.add_paragraph()
meta = doc.add_table(rows=1, cols=3); meta.alignment = WD_TABLE_ALIGNMENT.CENTER; line_borders(meta)
labels = [('VERSION', '1.0'), ('EFFECTIVE DATE', '26 June 2026'), ('REVIEW CYCLE', 'Annual')]
for i, (l, v) in enumerate(labels):
    cc = meta.cell(0, i); shade(cc, 'F8FAFC'); set_cell_margins(cc, 140, 140, 160, 160)
    pp = cc.paragraphs[0]; run(pp, l, size=8, bold=True, color=MUTED); pp.paragraph_format.space_after = Pt(2)
    pv = cc.add_paragraph(); run(pv, v, size=12, bold=True, color=NAVY)
doc.add_paragraph()
pf = doc.add_paragraph(); run(pf, 'Vyntra Property Services   ·   ABN 69 252 402 831   ·   Sydney, New South Wales   ·   CONFIDENTIAL', size=8.5, color=MUTED)
page_break()

# ============ WELCOME ============
kicker('Section 01 · Welcome aboard'); h1('Welcome to Vyntra')
para('Thank you for choosing to work with us. You are joining a network of trusted trade and cleaning professionals who help us deliver a premium, reliable service to property owners and managers across Sydney.', size=11, color=NAVY)
h2('Who we are')
para('Vyntra Property Services is a Sydney-based provider of premium property maintenance and cleaning. We connect property owners and managers with skilled subcontractors who get the job done properly, the first time.')
callout('Our mission', 'To make property care effortless for our clients by pairing them with dependable professionals and backing every job with clear standards, fast communication and fair, on-time payment.', fill='0F172A')
h2('What we expect')
checklist([
    'Quality — workmanship you would be proud to put your own name to.',
    'Reliability — turn up on time, communicate, and finish what you start.',
    'Respect — treat every client, property and colleague with courtesy.',
])
h2('Where you fit in')
para('As an independent subcontractor, you are the face of Vyntra on every job. The quality of your work and the way you treat our clients directly shapes our reputation — and your future opportunities with us.')
callout('Our commitment to you', 'In return, we commit to clear job briefs, responsive support, fair rates and reliable payment — so you can focus on doing great work.', fill='ECFDF3')
page_break()

# ============ HOW VYNTRA WORKS ============
kicker('Section 02 · The job lifecycle'); h1('How Vyntra works')
para('Every job follows the same simple flow — from the client’s first request through to your payment.', size=11, color=NAVY)
steps = [
    ('1', 'Client requests work', 'A property owner or manager books a job through Vyntra.'),
    ('2', 'Vyntra schedules the job', 'We confirm scope, timing and the agreed rate.'),
    ('3', 'Subcontractor receives job', 'You get the full brief, address and access details.'),
    ('4', 'Accept or decline', 'Confirm promptly so we can plan with confidence.'),
    ('5', 'Complete the work', 'Deliver to Vyntra’s quality and safety standards.'),
    ('6', 'Upload before & after photos', 'Required evidence for every job, no exceptions.'),
    ('7', 'Client approval', 'The client signs off that the work meets expectations.'),
    ('8', 'Invoice submitted', 'Send your tax invoice quoting the Vyntra job number.'),
    ('9', 'Payment processed', 'Paid to your nominated account on our payment terms.'),
]
for n, t, d in steps:
    p = doc.add_paragraph()
    run(p, n + '  ', size=12, bold=True, color=GOLD)
    run(p, t + '  —  ', size=10.5, bold=True, color=NAVY)
    run(p, d, size=9.5, color=MUTED)
    p.paragraph_format.space_after = Pt(4)
callout('Keep it moving', 'Fast acceptance, clean photos and a correct invoice are the three things that get you paid sooner.')
page_break()

# ============ CODE OF CONDUCT ============
kicker('Section 03 · Part 1 of 3'); h1('Code of Conduct — Presentation & people')
h2('Professional appearance')
checklist(['Wear clean, professional workwear and enclosed safety footwear.', 'Maintain good personal hygiene and a tidy, professional appearance.', 'Keep your tools, equipment and vehicle clean and well presented.', 'No Vyntra uniform is required — your own business branding is welcome.'])
h2('Customer service')
checklist(['Introduce yourself as working on behalf of Vyntra Property Services.', 'Greet clients politely and explain what you will be doing.', 'Listen, be patient, and never argue with a client.', 'Leave the client feeling looked after, not just serviced.'])
h2('Communication')
checklist(['Respond to Vyntra messages and calls promptly.', 'Flag delays, access issues or extra work straight away.', 'Keep all job communication professional and on record.', 'Never share your personal rates or details with clients.'])
h2('Punctuality')
checklist(['Arrive within the agreed appointment window.', 'Call ahead immediately if you are going to be late.', 'Allow enough time to complete the job properly.', 'Respect the client’s time as much as your own.'])
callout('If something goes wrong, tell us straight away', 'Problems and surprises happen on every job. Contact Vyntra immediately — you will never be in trouble for reporting a genuine issue early. Trying to hide a mistake is the only real mistake.', fill='ECFDF3')
page_break()

kicker('Section 03 · Part 2 of 3'); h1('Code of Conduct — On-site standards')
h2('Cleanliness')
checklist(['Use drop sheets and protect surfaces and floors.', 'Clean as you go and remove all rubbish and offcuts.', 'Leave the area cleaner than you found it.'])
h2('Before & after photos')
checklist(['Take clear before photos prior to starting.', 'Take matching after photos from the same angles.', 'Upload to the job before you leave site — every time.'])
h2('Property respect')
checklist(['Only access the areas required for the job.', 'Handle the client’s belongings with care.', 'Secure doors, windows and gates when you leave.', 'Report any pre-existing damage before you start.'])
h2('Privacy & confidentiality')
checklist(['Keep client details and access codes confidential.', 'Do not photograph clients or belongings beyond the work.', 'Never discuss one client’s job with another.'])
callout('Zero tolerance — smoking, alcohol & drugs', 'No smoking or vaping on or around any client property. Never attend a job under the influence of alcohol or drugs, and never consume them on site. Breaches result in immediate removal from the job and may end your engagement.', fill='FEF3F2')
page_break()

kicker('Section 03 · Part 3 of 3'); h1('Code of Conduct — Reputation & prohibited behaviour')
h2('Social media')
checklist(['Do not post client properties, addresses or people.', 'Get written approval before using Vyntra’s name or logo.', 'Never post anything that could embarrass a client or Vyntra.'])
h2('Representing Vyntra')
checklist(['Act honestly and ethically at all times.', 'Refer all quotes and bookings back to Vyntra.', 'Raise any client concern with Vyntra immediately.'])
h2('Prohibited behaviour — strictly not permitted')
checklist([
    'Soliciting clients for private or cash-in-hand work.',
    'Dishonesty, theft or misuse of client property.',
    'Aggressive, threatening or discriminatory conduct.',
    'Sharing client data or images without consent.',
    'Working while impaired by alcohol or drugs.',
    'Subcontracting a Vyntra job without our approval.',
    'Unsafe work practices or ignoring site safety.',
    'Falsifying photos, invoices or job records.',
], mark='✖')
callout('Consequences', 'Serious or repeated breaches of this Code may lead to suspension or immediate termination of your engagement, and where appropriate, referral to the relevant authorities.', fill='FEF3F2')
page_break()

# ============ PAYMENT POLICY ============
kicker('Section 04 · Part 1 of 2'); h1('Payment Policy — Getting paid')
para('We keep payments simple and predictable. Submit a correct tax invoice and we will pay it on time, every time.', size=11, color=NAVY)
h2('Payment schedule')
kv_table([
    ('Payment terms', '14 days from receipt of a valid tax invoice'),
    ('Payment method', 'Electronic funds transfer (EFT) to your nominated account'),
    ('Payment run', 'Processed weekly; cleared funds typically within 1–2 business days'),
    ('Remittance', 'Emailed to you each time a payment is made'),
], head=('Item', 'Standard'))
h2('Invoice requirements')
kv_table([
    ('Vyntra job number', 'Quoted on every invoice so we can match it quickly'),
    ('Business name & ABN', 'Must match the details we hold on file'),
    ('GST', 'If registered, show GST separately and the words "Tax Invoice"'),
    ('Description & date', 'Clear line items, site address and date completed'),
    ('Bank details', 'BSB and account number for EFT payment'),
], head=('Your tax invoice must include', 'Detail'))
callout('GST & ABN', 'Quote a valid ABN on every invoice. If you are registered for GST, add 10% GST and label it clearly. No ABN may mean we must withhold tax at the top rate.')
page_break()

kicker('Section 04 · Part 2 of 2'); h1('Payment Policy — Rates, call-outs & approvals')
h2('Rates & call-outs')
kv_table([
    ('Minimum call-out', 'A minimum charge applies to every attendance, as agreed for the job'),
    ('Standard hours', 'Mon–Fri, 7:00am–5:00pm at the agreed standard rate'),
    ('After-hours work', 'Evenings & weekends attract a higher agreed rate — confirm before attending'),
    ('Emergency work', 'Urgent call-outs are paid at the agreed emergency rate and prioritised'),
], head=('Type', 'How it works'))
h2('Approvals & changes')
kv_table([
    ('Extra work approval', 'Stop and get written approval from Vyntra before any work beyond the original scope. Unapproved extras may not be paid.'),
    ('Cancellation', 'If a client cancels with less than 24 hours’ notice after you are dispatched, a cancellation fee may apply as agreed.'),
    ('Variations', 'Any change to price, scope or timing must be confirmed by Vyntra in writing.'),
], head=('Situation', 'What to do'))
callout('Client non-payment', 'You are engaged and paid by Vyntra, not the client. Provided you have completed the work to standard and invoiced correctly, you are paid on our standard terms regardless of whether the client has paid us. Never chase payment directly from a client.')
page_break()

# ============ REQUIRED DOCUMENTS ============
kicker('Section 05 · Before your first job'); h1('Required documents')
para('Please provide the following so we can verify you and activate your account.', size=11, color=NAVY)
docs_list = [
    ('ABN — Australian Business Number', 'Mandatory'),
    ('Public Liability Insurance — current certificate, minimum $20 million', 'Mandatory'),
    ('Driver Licence — front & back', 'Mandatory'),
    ('White Card — construction induction', 'If applicable'),
    ('Police Check — national criminal history', 'If required'),
    ('Working With Children Check', 'If required'),
    ('Bank Details — BSB & account number', 'Mandatory'),
    ('Emergency Contact — name & phone', 'Mandatory'),
]
t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER; line_borders(t)
hd = t.rows[0].cells
for i, htxt in enumerate(('☐', 'Document', 'Requirement')):
    shade(hd[i], '0F172A'); set_cell_margins(hd[i])
    run(hd[i].paragraphs[0], htxt, size=9, bold=True, color=WHITE)
for ri, (d, r) in enumerate(docs_list):
    cells = t.add_row().cells
    if ri % 2 == 1:
        for cc in cells: shade(cc, 'F8FAFC')
    for cc in cells: set_cell_margins(cc)
    run(cells[0].paragraphs[0], '☐', size=13, color=NAVY)
    run(cells[1].paragraphs[0], d, size=9.5, color=NAVY, bold=True)
    run(cells[2].paragraphs[0], r, size=9.5, color=BODY)
for row in t.rows:
    row.cells[0].width = Cm(1.2); row.cells[1].width = Cm(10.6); row.cells[2].width = Cm(5.2)
doc.add_paragraph()
callout('How to submit', 'Attach clear PDF or photo copies to info@vyntrapropertyservices.com or upload them through your subcontractor portal. Keep insurance and checks current — expired documents will pause your jobs until renewed.')
page_break()

# ============ AGREEMENT ============
clauses = [
    ('1', 'Independent contractor relationship', 'You are engaged as an independent contractor, not an employee, partner or agent of Vyntra. You are responsible for your own tax, superannuation, insurances and leave. Nothing in this Agreement creates an employment relationship, and you may work for others provided you meet your obligations to us.'),
    ('2', 'Scope of work', 'Vyntra will offer you property maintenance and cleaning jobs from time to time. Each job will specify the work, location, timing and rate. You agree to perform accepted work with due care, skill and to the standards set out in this pack and any reasonable directions from Vyntra.'),
    ('3', 'Acceptance of jobs', 'You are free to accept or decline any job offered. Once you accept a job, you commit to completing it on time and to standard, or to giving Vyntra prompt notice if you genuinely cannot, so we can arrange an alternative.'),
    ('4', 'Your own equipment', 'Unless agreed otherwise, you supply your own tools, equipment, materials and transport, and are responsible for their condition, safety and maintenance.'),
    ('5', 'Insurance requirements', 'You must hold and maintain current Public Liability Insurance of at least $20 million (or another minimum Vyntra notifies you of), plus any insurance required by law for your trade — for example Workers Compensation where applicable. You agree to provide certificates of currency on request and to keep them current at all times.'),
    ('6', 'Work standards & safety', 'You agree to perform all work safely, lawfully and to a professional standard, complying with all work health and safety laws and the standards in this pack. You will take reasonable care for your own safety and that of others, and stop work and notify Vyntra if a situation is unsafe.'),
    ('7', 'Confidentiality', 'You will keep confidential all non-public information you learn through your engagement — including Vyntra’s clients, pricing, methods and systems — and use it only to perform the work. This obligation continues after your engagement ends.'),
    ('8', 'Privacy', 'You will handle any personal information of clients and occupants in line with the Privacy Act 1988 (Cth) and Vyntra’s directions — collecting only what is needed, keeping it secure, and never disclosing or using it for any purpose other than the job.'),
    ('9', 'Non-solicitation of Vyntra clients', 'During your engagement and for 12 months after it ends, you will not directly or indirectly solicit, accept or perform private work for any Vyntra client you were introduced to or serviced through Vyntra, without Vyntra’s written consent. All client opportunities must be referred back to Vyntra.'),
    ('10', 'Intellectual property', 'Any materials, photos, reports or records you create in performing the work — including before and after photos — are Vyntra’s property and may be used by Vyntra for service delivery, quality and marketing. You grant Vyntra a perpetual licence to use them for these purposes.'),
    ('11', 'Liability & indemnity', 'You are responsible for the work you perform and for any loss, damage or claim arising from your acts, omissions or negligence. To the extent permitted by law, you indemnify Vyntra against such claims. Nothing in this clause limits rights that cannot be excluded under the Australian Consumer Law.'),
    ('12', 'Termination', 'Either party may end this Agreement at any time by giving reasonable notice in writing. Vyntra may suspend or end your engagement immediately for a serious breach — including a breach of the Code of Conduct, safety, confidentiality or non-solicitation terms. Ending the Agreement does not affect rights or payments accrued beforehand.'),
    ('13', 'Subcontracting & assignment', 'You will not subcontract, assign or transfer a Vyntra job to anyone else without Vyntra’s prior written consent. Accepted jobs are personal to you and the team we have approved.'),
    ('14', 'Dispute resolution', 'If a dispute arises, both parties agree to first raise it in good faith and try to resolve it through discussion. If it cannot be resolved within a reasonable time, the parties may refer the matter to mediation before commencing legal proceedings, except where urgent relief is needed.'),
    ('15', 'Governing law', 'This Agreement is governed by the laws of New South Wales, Australia. The parties submit to the non-exclusive jurisdiction of the courts of New South Wales.'),
    ('16', 'Entire agreement', 'This Agreement, together with the Code of Conduct, Payment Policy and each job’s brief, forms the entire agreement between us. If any part is found to be invalid, the rest continues to apply. Vyntra may update this pack from time to time and will notify you of material changes.'),
]
kicker('Section 06 · Independent Contractor Agreement'); h1('Independent Contractor Agreement')
para('This Agreement sets out the terms on which you provide services to Vyntra as an independent contractor. It is written in plain English — please read it before you sign the declaration.', size=11, color=NAVY)
kv_table([
    ('Principal', 'Vyntra Property Services'),
    ('ABN', '69 252 402 831'),
    ('Contractor', 'As named on the signed declaration'),
    ('Governing law', 'New South Wales, Australia'),
], head=('Party', 'Detail'))
for i, (n, t, d) in enumerate(clauses):
    p = doc.add_paragraph()
    run(p, n + '.  ', size=10.5, bold=True, color=GOLD)
    run(p, t, size=10.5, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(6)
    pb = doc.add_paragraph(); run(pb, d, size=9.5)
    pb.paragraph_format.space_after = Pt(4)
    if i == 7:  # break agreement across two pages
        page_break()
        kicker('Section 06 · Independent Contractor Agreement (continued)'); h1('Agreement — continued')
callout('Plain-English summary', 'Be safe and insured, do great work, look after our clients and their information, refer client opportunities back to us, and we will keep the jobs and payments flowing. This document is a clear summary of our working arrangement and is not legal advice.', fill='ECFDF3')
page_break()

# ============ FAQ ============
kicker('Section 07 · Quick answers'); h1('Frequently asked questions')
faqs = [
    ('When do I get paid?', 'Within 14 days of us receiving a correct tax invoice. Payments go out by EFT in our weekly run.'),
    ('Can I decline jobs?', 'Yes. You choose which jobs to accept. Please respond quickly so we can offer it to someone else if needed.'),
    ('What if I am running late?', 'Call ahead as soon as you know. Keep the client and Vyntra informed — a quick heads-up prevents most problems.'),
    ('What if the customer wants additional work?', 'Do not start it. Get written approval from Vyntra first so the extra scope and price are agreed and you get paid for it.'),
    ('Who do I contact?', 'Your Vyntra coordinator on 0451 510 026 or info@vyntrapropertyservices.com for anything job-related.'),
    ('What happens if equipment is damaged?', 'Stop, make the area safe, and report it to Vyntra immediately with photos. Do not attempt to hide or quietly fix damage.'),
]
for q, a in faqs:
    p = doc.add_paragraph(); run(p, 'Q.  ', size=10.5, bold=True, color=GOLD); run(p, q, size=10.5, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(1)
    pa = doc.add_paragraph(); run(pa, a, size=9.5)
    pa.paragraph_format.left_indent = Cm(0.55); pa.paragraph_format.space_after = Pt(3)
callout('Still have a question?', 'Reach the Vyntra team on 0451 510 026 or email info@vyntrapropertyservices.com. We are here to help you do your best work.')
page_break()

# ============ DECLARATION ============
kicker('Section 08 · Read · agree · sign'); h1('Declaration & agreement')
para('By signing below, I confirm that I have read, understood and agree to be bound by the following. This page can be completed and signed electronically, or printed and signed by hand.', size=11, color=NAVY)
checklist([
    'I agree to the Independent Contractor Agreement',
    'I agree to the Code of Conduct',
    'I agree to the Payment Policy',
    'I agree to the Safety Requirements',
    'I agree to maintain Confidentiality',
    'I agree to protect Client Privacy',
    'I agree to the Non-Solicitation Policy',
], mark='☐')
p_acc = doc.add_paragraph()
run(p_acc, '☐  ', size=13, color=NAVY)
run(p_acc, 'I confirm that all information and documents I have provided are accurate and current, and I understand that providing false or expired documents may result in suspension or removal from the Vyntra network.', size=9.5)
p_acc.paragraph_format.left_indent = Cm(0.4); p_acc.paragraph_format.space_before = Pt(4); p_acc.paragraph_format.space_after = Pt(4)
h2('Your details')
fields = [('Full Name', ''), ('ABN', ''), ('Company Name', ''), ('Mobile Number', ''), ('Email', ''), ('Date', '')]
ft = doc.add_table(rows=3, cols=2); ft.alignment = WD_TABLE_ALIGNMENT.CENTER; line_borders(ft)
flat = [fields[0], fields[1], fields[2], fields[3], fields[4], fields[5]]
idx = 0
for r in range(3):
    for c in range(2):
        cell = ft.cell(r, c); set_cell_margins(cell, 120, 200, 160, 160)
        lbl, _ = flat[idx]; idx += 1
        run(cell.paragraphs[0], lbl.upper(), size=8, bold=True, color=MUTED)
        cell.add_paragraph().add_run(' ')
for row in ft.rows:
    row.cells[0].width = Cm(8.5); row.cells[1].width = Cm(8.5)
doc.add_paragraph()
st = doc.add_table(rows=1, cols=1); st.alignment = WD_TABLE_ALIGNMENT.CENTER; line_borders(st, '94A3B8')
sc = st.cell(0, 0); shade(sc, 'F8FAFC'); set_cell_margins(sc, 200, 700, 160, 160)
run(sc.paragraphs[0], 'SIGNATURE', size=8, bold=True, color=MUTED)
doc.add_paragraph()
callout('Return this page', 'Send the completed declaration to info@vyntrapropertyservices.com with your documents to activate your account.', fill='ECFDF3')
pf2 = doc.add_paragraph(); run(pf2, 'Vyntra Property Services · ABN 69 252 402 831 · Level 12, 1 Sydney Avenue, Sydney NSW · This declaration forms part of your engagement with Vyntra.', size=8, color=MUTED)
page_break()

# ============ WHAT HAPPENS NEXT ============
kicker('Section 09 · After you sign'); h1('What happens next?')
para('Once you have signed your declaration and sent it back with your documents, here is how we will get you ready for your first job.', size=11, color=NAVY)
next_steps = [
    ('1', 'You submit', 'Return your signed declaration and required documents to Vyntra.'),
    ('2', 'We review & verify', 'We check your documents, insurance and any required clearances — usually within 2–3 business days.'),
    ('3', 'Approval', 'Once everything checks out, we confirm that you are approved to work with Vyntra.'),
    ('4', 'Account activation', 'We set up your subcontractor profile and add you to our active network.'),
    ('5', 'Your first jobs', 'You start receiving job offers that match your trade, availability and service area.'),
]
for n, t, d in next_steps:
    p = doc.add_paragraph()
    run(p, n + '  ', size=12, bold=True, color=GOLD)
    run(p, t + '  —  ', size=10.5, bold=True, color=NAVY)
    run(p, d, size=9.5, color=MUTED)
    p.paragraph_format.space_after = Pt(4)
callout('If anything is missing', 'If anything needs updating, we will let you know exactly what is required — so nothing holds up your start.', fill='ECFDF3')
h2('Contact Vyntra')
ct = doc.add_table(rows=1, cols=3); ct.alignment = WD_TABLE_ALIGNMENT.CENTER; no_borders(ct)
contacts = [('PHONE', '0451 510 026'), ('EMAIL', 'info@vyntrapropertyservices.com'), ('AFTER-HOURS & EMERGENCIES', '0451 510 026')]
for i, (l, v) in enumerate(contacts):
    cc = ct.cell(0, i); shade(cc, '0F172A'); set_cell_margins(cc, 160, 160, 180, 180)
    run(cc.paragraphs[0], l, size=8, bold=True, color=GOLD_BR); cc.paragraphs[0].paragraph_format.space_after = Pt(2)
    pv = cc.add_paragraph(); run(pv, v, size=11, bold=True, color=WHITE)
for i in range(3):
    ct.cell(0, i).width = Cm(5.67)
doc.add_paragraph()
callout('Welcome to the Vyntra network', 'We are glad to have you on board. Do great work, keep us in the loop, and we will keep the jobs coming.', fill='0F172A')

out = 'dist/Vyntra-Subcontractor-Onboarding-Pack.docx'
doc.save(out)
print('Saved', out)
